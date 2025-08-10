import json
import logging
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponseNotFound, HttpResponse
from django.template.loader import render_to_string
from django.views.decorators.csrf import csrf_exempt, csrf_protect
from django.core.exceptions import ValidationError
from django.conf import settings
from django.contrib.auth import login
from .ai_services import EnhancedAICoverLetterService
from .views_upload_cv_optimized import upload_cv_optimized
from .views_upload_cv_analyzer import upload_cv_analyzer
from .models import AICoverLetter, CVAnalysis, CV, CVSection, UploadedCV
from .enhanced_forms import EnhancedAICoverLetterForm
from .forms import CVCreationForm
from django.contrib import messages

logger = logging.getLogger(__name__)

def load_cv_template(request, template_name):
    """Loads a CV template for editing or preview"""
    preview_mode = request.GET.get('preview', False)
    
    try:
        # First try the templates directory
        template_path = f'builder/templates/builder/templates/{template_name}.html'
        context = {
            'cv': {},
            'preview_mode': preview_mode
        }
        html = render_to_string(template_path, context)
        return HttpResponse(html)
    except Exception as e:
        logger.error(f"Error loading template {template_name}: {str(e)}")
        return HttpResponse('Template not found', status=404)

def templates(request):
    """Display available CV templates"""
    templates = [
        {
            'id': 'basic',
            'name': 'Basic Template',
            'description': 'A clean and professional CV template suitable for most industries',
            'preview_image': 'builder/images/templates/basic-preview.png'
        },
        # Add more templates as needed
    ]
    return render(request, 'builder/templates.html', {'templates': templates})

def save_cv_content(request):
    """Saves CV content from the editor"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            # Here you would save the CV data to your database
            # For now, we'll just return success
            return JsonResponse({'status': 'success'})
        except Exception as e:
            logger.error(f"Error saving CV content: {str(e)}")
            return JsonResponse({'status': 'error', 'message': str(e)})
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'})

@login_required
def enhanced_ai_cover_letter(request):
    """Enhanced AI cover letter generator view"""
    if request.method == 'POST':
        form = EnhancedAICoverLetterForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                # Process form data
                job_title = form.cleaned_data['job_title']
                job_description = form.cleaned_data['job_description']
                tone = form.cleaned_data.get('tone', 'professional')
                template_type = form.cleaned_data.get('template_type', 'standard')
                
                # Get CV text
                cv_text = form.cleaned_data.get('cv_text', '')
                if form.cleaned_data.get('uploaded_cv'):
                    # Process uploaded CV file
                    cv_file = form.cleaned_data['uploaded_cv']
                    try:
                        if cv_file.content_type.startswith('text'):
                            cv_text = cv_file.read().decode('utf-8', errors='ignore')[:5000]
                        else:
                            cv_text = f"Binary file uploaded: {cv_file.name}"
                    except Exception as e:
                        cv_text = f"Error reading file: {str(e)}"
                
                # Generate cover letter
                service = EnhancedAICoverLetterService()
                try:
                    if not service.client:
                        logger.error("OpenAI API key not configured")
                        messages.error(request, "OpenAI API is not properly configured. Please try again later or contact support.")
                        return render(request, 'builder/enhanced_ai_cover_letter.html', {'form': form})

                    cv_insights = service.extract_cv_insights(cv_text)
                    if not cv_insights:
                        logger.error("Failed to extract CV insights")
                        messages.error(request, "Failed to analyze CV. Please try again.")
                        return render(request, 'builder/enhanced_ai_cover_letter.html', {'form': form})

                    job_match = service.match_cv_to_job(cv_insights, job_title, job_description)
                    if not job_match:
                        logger.error("Failed to match CV to job")
                        messages.error(request, "Failed to match CV with job requirements. Please try again.")
                        return render(request, 'builder/enhanced_ai_cover_letter.html', {'form': form})

                    cover_letter = service.generate_tailored_cover_letter(
                        cv_insights, job_match, job_title, job_description, tone, template_type
                    )
                
                    # Save to database
                    ai_cover_letter = AICoverLetter.objects.create(
                        user=request.user,
                        job_title=job_title,
                        job_description=job_description,
                        generated_letter=cover_letter,
                        cv_analysis=cv_text[:500],
                        tone=tone,
                        template_type=template_type
                    )
                except Exception as e:
                    logger.error(f"Error generating cover letter: {str(e)}")
                    messages.error(request, "An error occurred while generating your cover letter. Please try again.")
                    return render(request, 'builder/enhanced_ai_cover_letter.html', {'form': form})
                
                return render(request, 'builder/enhanced_ai_cover_letter_result.html', {
                    'cover_letter': cover_letter,
                    'cv_insights': cv_insights,
                    'job_match': job_match,
                    'ai_cover_letter': ai_cover_letter
                })
                
            except Exception as e:
                logger.error(f"Cover letter generation failed: {str(e)}")
                import traceback
                logger.error(f"Full traceback: {traceback.format_exc()}")
                
                # Provide more specific error message
                error_message = 'Failed to generate cover letter. Please try again.'
                if 'Job title and description are required' in str(e):
                    error_message = 'Please provide both job title and job description.'
                elif 'OpenAI' in str(e):
                    error_message = 'AI service temporarily unavailable. Please try again later.'
                
                return render(request, 'builder/enhanced_ai_cover_letter.html', {
                    'form': form,
                    'error': error_message
                })
    else:
        # Check if CV ID is provided in URL
        cv_id = request.GET.get('cv_id')
        initial_data = {}
        
        if cv_id:
            try:
                uploaded_cv = UploadedCV.objects.get(id=cv_id, user=request.user)
                initial_data = {
                    'uploaded_cv': uploaded_cv,
                    'cv_text': uploaded_cv.extracted_text
                }
            except UploadedCV.DoesNotExist:
                pass
        
        form = EnhancedAICoverLetterForm(initial=initial_data)
    
    return render(request, 'builder/enhanced_ai_cover_letter.html', {'form': form})

@login_required
@csrf_protect
def ajax_generate_cover_letter(request):
    """AJAX endpoint for real-time cover letter generation"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Invalid request method'}, status=405)
    
    try:
        data = json.loads(request.body)
        job_title = data.get('job_title')
        job_description = data.get('job_description')
        cv_text = data.get('cv_text', '')
        tone = data.get('tone', 'professional')
        
        if not all([job_title, job_description]):
            return JsonResponse({'error': 'Missing required fields'}, status=400)
        
        logger.info(f"Starting AJAX cover letter generation for job: {job_title}")
        service = EnhancedAICoverLetterService()
        logger.info("AI service initialized for AJAX request")
        
        cv_insights = service.extract_cv_insights(cv_text)
        logger.info(f"CV insights extracted: {len(cv_insights.get('skills', []))} skills found")
        
        job_match = service.match_cv_to_job(cv_insights, job_title, job_description)
        logger.info("Job matching completed")
        
        cover_letter = service.generate_tailored_cover_letter(
            cv_insights, job_match, job_title, job_description, tone
        )
        logger.info(f"Cover letter generated: {len(cover_letter)} characters")
        
        return JsonResponse({
            'success': True,
            'cover_letter': cover_letter,
            'cv_insights': cv_insights,
            'job_match': job_match
        })
        
    except Exception as e:
        logger.error(f"AJAX generation failed: {str(e)}")
        return JsonResponse({'error': 'Generation failed'}, status=500)

@login_required
def edit_generated_letter(request, pk):
    """Edit saved generated cover letter"""
    ai_cover_letter = get_object_or_404(AICoverLetter, pk=pk, user=request.user)
    
    if request.method == 'POST':
        new_letter = request.POST.get('generated_letter', '')
        if new_letter:
            ai_cover_letter.generated_letter = new_letter
            ai_cover_letter.save()
            return redirect('enhanced_ai_cover_letter')
    
    return render(request, 'builder/edit_generated_letter.html', {
        'ai_cover_letter': ai_cover_letter
    })

@login_required
def cv_analysis_detail(request, pk):
    """Detailed CV analysis view"""
    cv_analysis = get_object_or_404(CVAnalysis, pk=pk, user=request.user)
    return render(request, 'builder/cv_analysis_detail.html', {
        'cv_analysis': cv_analysis
    })

def register(request):
    """User registration view"""
    from .enhanced_forms import CustomUserCreationForm
    
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('builder:home')
    else:
        form = CustomUserCreationForm()
    
    return render(request, 'registration/register.html', {'form': form})

def home(request):
    """Home page view - upload only"""
    return render(request, 'builder/home_upload_only.html')

@login_required
def dashboard(request):
    """Dashboard view - upload only"""
    uploaded_cvs = UploadedCV.objects.select_related('user').filter(user=request.user).order_by('-uploaded_at')
    ai_cover_letters = AICoverLetter.objects.select_related('uploaded_cv__user').filter(uploaded_cv__user=request.user).order_by('-created_at')
    
    return render(request, 'builder/dashboard_upload_only.html', {
        'uploaded_cvs': uploaded_cvs,
        'ai_cover_letters': ai_cover_letters
    })

def template_previews(request):
    """View for CV template previews gallery"""
    try:
        logger.info("Accessing template preview gallery")
        return render(request, 'builder/template_previews.html')
    except Exception as e:
        logger.error(f"Error rendering template preview gallery: {str(e)}")
        messages.error(request, "Unable to load template preview gallery. Please try again later.")
        return redirect('builder:home')

def template_preview(request, template_name):
    """View for individual CV template preview"""
    template_names = ['modern', 'classic', 'minimal', 'creative']
    try:
        if template_name not in template_names:
            logger.warning(f"Invalid template name requested: {template_name}")
            messages.warning(request, f"Template '{template_name}' not found.")
            return HttpResponseNotFound('Template not found')
        
        logger.info(f"Accessing preview for template: {template_name}")
        return render(request, f'builder/templates/{template_name}_preview.html')
    except Exception as e:
        logger.error(f"Error rendering template preview for {template_name}: {str(e)}")
        messages.error(request, f"Unable to load template preview. Please try again later.")
        return redirect('builder:template_previews')

def create_cv(request):
    """Create CV view with form handling"""
    if request.method == 'POST':
        form = CVCreationForm(request.POST)
        
        if form.is_valid():
            try:
                # Create new CV instance
                cv = CV.objects.create(
                    title=f"{form.cleaned_data['full_name']} - {form.cleaned_data['title']}",
                    user=request.user
                )
                
                # Create CV sections
                # Personal Information Section
                personal_info = f"""
                {form.cleaned_data['full_name']}
                {form.cleaned_data['title']}
                {form.cleaned_data['email']} | {form.cleaned_data['phone']} | {form.cleaned_data['location']}
                {form.cleaned_data['linkedin'] if form.cleaned_data['linkedin'] else ''}
                """
                
                CVSection.objects.create(
                    cv=cv,
                    content=personal_info
                )
                
                # Summary Section
                if form.cleaned_data['summary']:
                    CVSection.objects.create(
                        cv=cv,
                        content=f"Professional Summary: {form.cleaned_data['summary']}"
                    )
                
                # Work Experience Section
                if form.experience_data:
                    experience_content = "Work Experience:\n"
                    for exp in form.experience_data:
                        experience_content += f"""
                        {exp['title']} at {exp['company']}
                        {exp['duration']} | {exp['location']}
                        {exp['description']}
                        """
                    CVSection.objects.create(
                        cv=cv,
                        content=experience_content
                    )
                
                # Education Section
                if form.education_data:
                    education_content = "Education:\n"
                    for edu in form.education_data:
                        education_content += f"""
                        {edu['degree']} - {edu['school']}
                        {edu['year']} | {edu['location']}
                        """
                    CVSection.objects.create(
                        cv=cv,
                        content=education_content
                    )
                
                # Skills Section
                if form.cleaned_data['skills']:
                    CVSection.objects.create(
                        cv=cv,
                        content=f"Skills: {form.cleaned_data['skills']}"
                    )
                
                messages.success(request, 'CV created successfully!')
                return redirect('cv_detail', pk=cv.pk)
                
            except Exception as e:
                logger.error(f"CV creation failed: {str(e)}")
                messages.error(request, 'Failed to create CV. Please try again.')
                
    else:
        form = CVCreationForm()
    
    return render(request, 'builder/create_cv.html', {'form': form})

def cv_detail(request, pk):
    """CV detail view"""
    cv = get_object_or_404(CV, pk=pk, user=request.user)
    sections = CVSection.objects.filter(cv=cv)
    
    # Parse sections into structured data
    cv_data = {
        'personal_info': {},
        'summary': '',
        'experience': [],
        'education': [],
        'skills': []
    }
    
    for section in sections:
        content = section.content.strip()
        if content.startswith('Professional Summary:'):
            cv_data['summary'] = content.replace('Professional Summary:', '').strip()
        elif content.startswith('Skills:'):
            skills_text = content.replace('Skills:', '').strip()
            cv_data['skills'] = [skill.strip() for skill in skills_text.split(',') if skill.strip()]
        elif content.startswith('Work Experience:'):
            # Parse work experience
            lines = content.split('\n')
            current_exp = {}
            for line in lines:
                line = line.strip()
                if line and 'at' in line.lower() and not line.startswith('Work Experience:'):
                    # This is a job title line
                    parts = line.split(' at ')
                    if len(parts) >= 2:
                        current_exp = {
                            'title': parts[0].strip(),
                            'company': parts[1].strip()
                        }
                elif '•' in line and current_exp:
                    # This is duration and location
                    parts = line.split('•')
                    if len(parts) >= 2:
                        current_exp['duration'] = parts[0].strip()
                        current_exp['location'] = parts[1].strip()
                elif line and current_exp and not line.startswith('Work Experience:'):
                    # This is description
                    if 'description' not in current_exp:
                        current_exp['description'] = line
                    else:
                        current_exp['description'] += ' ' + line
                    if current_exp not in cv_data['experience']:
                        cv_data['experience'].append(current_exp.copy())
        elif content.startswith('Education:'):
            # Parse education
            lines = content.split('\n')
            current_edu = {}
            for line in lines:
                line = line.strip()
                if line and '-' in line and not line.startswith('Education:'):
                    # This is degree and school
                    parts = line.split(' - ')
                    if len(parts) >= 2:
                        current_edu = {
                            'degree': parts[0].strip(),
                            'school': parts[1].strip()
                        }
                elif '•' in line and current_edu:
                    # This is year and location
                    parts = line.split('•')
                    if len(parts) >= 2:
                        current_edu['year'] = parts[0].strip()
                        current_edu['location'] = parts[1].strip()
                        if current_edu not in cv_data['education']:
                            cv_data['education'].append(current_edu.copy())
        else:
            # Parse personal info
            lines = content.split('\n')
            if len(lines) >= 3:
                cv_data['personal_info']['name'] = lines[0].strip()
                cv_data['personal_info']['title'] = lines[1].strip()
                contact_parts = lines[2].split('|')
                if len(contact_parts) >= 3:
                    cv_data['personal_info']['email'] = contact_parts[0].strip()
                    cv_data['personal_info']['phone'] = contact_parts[1].strip()
                    cv_data['personal_info']['location'] = contact_parts[2].strip()
                if len(lines) > 3 and lines[3].strip():
                    cv_data['personal_info']['linkedin'] = lines[3].strip()
    
    return render(request, 'builder/cv_detail.html', {
        'cv': cv,
        'cv_data': cv_data
    })

def add_cv_section(request, pk):
    """Add CV section view"""
    return render(request, 'builder/add_cv_section.html', {'pk': pk})

def edit_cv_section(request, pk):
    """Edit CV section view"""
    return render(request, 'builder/edit_cv_section.html', {'pk': pk})

def delete_cv_section(request, pk):
    """Delete CV section view"""
    return render(request, 'builder/delete_cv_section.html', {'pk': pk})

def generate_cover_letter(request):
    """Generate cover letter view - now uses enhanced AI"""
    return enhanced_ai_cover_letter(request)

@login_required
def upload_cv(request):
    """Upload CV view with file upload functionality"""
    from .forms import UploadedCVForm
    
    if request.method == 'POST':
        form = UploadedCVForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                uploaded_cv = form.save(commit=False)
                uploaded_cv.user = request.user
                
                # Process the uploaded file
                uploaded_file = request.FILES['file']
                uploaded_cv.original_filename = uploaded_file.name
                
                # Validate file security
                from .file_validators import FileValidator
                try:
                    FileValidator.validate_file(uploaded_file)
                    actual_mime_type = FileValidator.get_file_type(uploaded_file)
                except ValidationError as e:
                    messages.error(request, f"File validation failed: {str(e)}")
                    return redirect('builder:upload_cv')
                
                # Extract text from PDF/DOCX with proper error handling
                try:
                    if actual_mime_type == 'application/pdf':
                        try:
                            import PyPDF2
                            pdf_reader = PyPDF2.PdfReader(uploaded_file)
                            text = ""
                            for page in pdf_reader.pages:
                                text += page.extract_text()
                            uploaded_cv.extracted_text = text[:5000] if text.strip() else "No text found in PDF"
                        except Exception as pdf_error:
                            logger.error(f"PDF extraction failed: {str(pdf_error)}")
                            uploaded_cv.extracted_text = "PDF text extraction failed"
                            
                    elif actual_mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                        try:
                            import docx
                            doc = docx.Document(uploaded_file)
                            text = ""
                            for paragraph in doc.paragraphs:
                                text += paragraph.text + "\n"
                            uploaded_cv.extracted_text = text[:5000] if text.strip() else "No text found in document"
                        except Exception as docx_error:
                            logger.error(f"DOCX extraction failed: {str(docx_error)}")
                            uploaded_cv.extracted_text = "Document text extraction failed"
                    else:
                        uploaded_cv.extracted_text = "Text extraction not supported for this format"
                        
                    uploaded_cv.processed = True
                    uploaded_cv.save()
                    
                    # Create CV analysis
                    analysis = {
                        'overall_score': 78,
                        'strengths': [
                            'Strong technical skills',
                            'Clear career progression',
                            'Quantifiable achievements',
                            'Relevant experience',
                            'Professional formatting'
                        ],
                        'improvements': [
                            'Add more keywords',
                            'Include metrics',
                            'Optimize for ATS',
                            'Add summary section'
                        ],
                        'keywords': {
                            'present': ['Python', 'Django', 'JavaScript', 'SQL'],
                            'missing': ['React', 'AWS', 'Docker'],
                            'suggestions': ['Leadership', 'Communication', 'Problem-solving']
                        },
                        'experience_level': 'Mid-level',
                        'ats_compatibility': 85
                    }
                    
                    cv_analysis = CVAnalysis.objects.create(
                        uploaded_cv=uploaded_cv,
                        overall_score=analysis.get('overall_score', 0),
                        strengths=analysis.get('strengths', []),
                        improvements=analysis.get('improvements', []),
                        keywords=analysis.get('keywords', {}),
                        experience_level=analysis.get('experience_level', 'Unknown'),
                        ats_compatibility=analysis.get('ats_compatibility', 0)
                    )
                    
                    return render(request, 'builder/upload_cv_success.html', {
                        'uploaded_cv': uploaded_cv,
                        'cv_analysis': cv_analysis
                    })
                    
                except Exception as e:
                    logger.error(f"File processing failed: {str(e)}")
                    uploaded_cv.extracted_text = f"Error processing file: {str(e)}"
                    uploaded_cv.save()
                    
            except Exception as e:
                logger.error(f"Upload failed: {str(e)}")
                return render(request, 'builder/upload_cv.html', {
                    'form': form,
                    'error': 'Upload failed. Please try again.'
                })
    else:
        form = UploadedCVForm()
    
    return render(request, 'builder/upload_cv.html', {'form': form})

def ai_cover_letter(request):
    """AI cover letter view"""
    return render(request, 'builder/ai_cover_letter.html')

# CV analyzer view removed as part of CV analyzer functionality removal

def templates(request):
    """Templates view"""
    return render(request, 'builder/templates.html')

def cover_letter_templates(request):
    """Cover letter templates editor view"""
    return render(request, 'builder/cover_letter_templates.html')

@login_required
def cv_analyzer(request):
    """CV Analyzer view with AI-powered analysis"""
    if request.method == 'POST':
        try:
            cv_file = request.FILES.get('cv_file')
            if not cv_file:
                messages.error(request, 'Please upload a CV file.')
                return render(request, 'builder/cv_analyzer.html')
            
            # Validate file
            from .file_validators import FileValidator
            try:
                FileValidator.validate_file(cv_file)
                actual_mime_type = FileValidator.get_file_type(cv_file)
            except ValidationError as e:
                messages.error(request, f"File validation failed: {str(e)}")
                return render(request, 'builder/cv_analyzer.html')
            
            # Extract text from CV
            cv_text = ""
            try:
                if actual_mime_type == 'application/pdf':
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(cv_file)
                    for page in pdf_reader.pages:
                        cv_text += page.extract_text()
                elif actual_mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                    import docx
                    doc = docx.Document(cv_file)
                    for paragraph in doc.paragraphs:
                        cv_text += paragraph.text + "\n"
                else:
                    messages.error(request, 'Unsupported file format.')
                    return render(request, 'builder/cv_analyzer.html')
                    
                if not cv_text.strip():
                    messages.error(request, 'No text found in the uploaded file.')
                    return render(request, 'builder/cv_analyzer.html')
                    
            except Exception as e:
                logger.error(f"Text extraction failed: {str(e)}")
                messages.error(request, 'Failed to extract text from the CV. Please try a different file.')
                return render(request, 'builder/cv_analyzer.html')
            
            # Analyze CV using AI service
            logger.info("Starting CV analysis...")
            service = EnhancedAICoverLetterService()
            
            # Get comprehensive analysis
            analysis_data = service.analyze_cv_comprehensive(cv_text)
            logger.info("CV analysis completed successfully")
            
            return render(request, 'builder/cv_analyzer.html', {
                'analysis': analysis_data,
                'cv_text': cv_text[:500] + '...' if len(cv_text) > 500 else cv_text
            })
            
        except Exception as e:
            logger.error(f"CV analysis failed: {str(e)}")
            messages.error(request, 'Analysis failed. Please try again.')
            return render(request, 'builder/cv_analyzer.html')
    
    return render(request, 'builder/cv_analyzer.html')

def template_detail(request, pk):
    """Template detail view"""
    return render(request, 'builder/template_detail.html', {'pk': pk})

def create_template(request):
    """Create template view"""
    return render(request, 'builder/create_template.html')

def edit_template(request, pk):
    """Edit template view"""
    return render(request, 'builder/edit_template.html', {'pk': pk})

def delete_cv(request, pk):
    """Delete CV view"""
    cv = get_object_or_404(CV, pk=pk, user=request.user)
    
    if request.method == 'POST':
        try:
            cv.delete()
            messages.success(request, 'CV deleted successfully!')
            return redirect('dashboard')
        except Exception as e:
            logger.error(f"Failed to delete CV {pk}: {str(e)}")
            messages.error(request, 'Failed to delete CV. Please try again.')
            return redirect('dashboard')
    
    # For GET requests, redirect to dashboard since confirmation is handled by JS
    return redirect('dashboard')

def delete_uploaded_cv(request, pk):
    """Delete uploaded CV view"""
    uploaded_cv = get_object_or_404(UploadedCV, pk=pk, user=request.user)
    
    if request.method == 'POST':
        try:
            # Delete the file from storage
            if uploaded_cv.file:
                uploaded_cv.file.delete(save=False)
            uploaded_cv.delete()
            messages.success(request, 'Uploaded CV deleted successfully!')
            return redirect('dashboard')
        except Exception as e:
            logger.error(f"Failed to delete uploaded CV {pk}: {str(e)}")
            messages.error(request, 'Failed to delete uploaded CV. Please try again.')
            return redirect('dashboard')
    
    # For GET requests, redirect to dashboard since confirmation is handled by JS
    return redirect('dashboard')
